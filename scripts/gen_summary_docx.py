# -*- coding: utf-8 -*-
"""生成 DeepTutor 课程点播系统工作总结 Word 文档。"""
import sys
sys.path.insert(0, r"D:\Users\zzh\Desktop\DeepTutor项目\deep-tutor-plus\.venv\Lib\site-packages")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# 页面边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

# 默认字体（中文宋体 + 西文 Calibri）
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

def set_cn(run, font="宋体", size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def h1(text):
    p = doc.add_heading(level=1)
    r = p.add_run(text)
    set_cn(r, "黑体", 16, True, RGBColor(0x1F, 0x1D, 0x1B))
    return p

def h2(text):
    p = doc.add_heading(level=2)
    r = p.add_run(text)
    set_cn(r, "黑体", 13, True, RGBColor(0xB0, 0x50, 0x1E))
    return p

def h3(text):
    p = doc.add_heading(level=3)
    r = p.add_run(text)
    set_cn(r, "黑体", 11.5, True, RGBColor(0x1F, 0x1D, 0x1B))
    return p

def para(text, bold=False, size=11, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(text)
    set_cn(r, "宋体", size, bold)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.35
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_cn(r1, "宋体", 11, True)
    r2 = p.add_run(text)
    set_cn(r2, "宋体", 11)
    return p

def make_table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for j, htext in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.paragraphs[0].text = ""
        r = cell.paragraphs[0].add_run(htext)
        set_cn(r, "黑体", 10.5, True)
    # 数据
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.rows[i+1].cells[j]
            cell.paragraphs[0].text = ""
            r = cell.paragraphs[0].add_run(str(val))
            set_cn(r, "宋体", 10)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Cm(w)
    doc.add_paragraph()  # 表后空行
    return t

# ═══════════════════════ 封面 ═══════════════════════
for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DeepTutor 课程点播系统")
set_cn(r, "黑体", 26, True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("工作总结与阶段汇报")
set_cn(r, "黑体", 18, True, RGBColor(0xB0, 0x50, 0x1E))

for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("汇报人：研发组\n日期：2026 年 8 月")
set_cn(r, "宋体", 12)

doc.add_page_break()

# ═══════════════════════ 一、背景与目标 ═══════════════════════
h1("一、背景与目标")

h2("1.1 原始需求")
para("按照立项时确定的方案，我们希望把 DeepTutor 与 seerr 融合，做成一个课程点播系统："
     "在 DeepTutor 原有的 AI 辅导能力（对话、写作、协同等）之上，增加课程视频的入库、管理和点播能力，"
     "后续可以持续把课程视频上传到平台，供学生按需观看。")

h2("1.2 技术选型的调整：seerr → Jellyfin")
para("在动手前我做了技术验证，结论是：seerr 并不适合这个场景，需要换成 Jellyfin。", bold=True)
bullet("seerr 本质上是一个「点单」前端，它的数据源是 TMDB 的影视元数据（海报、演员、简介），"
       "必须绑定 TMDB 才能工作。而我们要上传的是自制课程视频，在 TMDB 上没有任何条目，"
       "seerr 对它完全无能为力。", bold_prefix="为什么 seerr 不行：")
bullet("Jellyfin 是完整的媒体服务器：支持文件入库、元数据扫描、转码、多端播放，"
       "并且不需要外部数据源——把课程视频按「剧集」结构（SxxExx 命名）放进去就能识别。"
       "它天然适合做课程点播的后端。", bold_prefix="为什么选 Jellyfin：")
bullet("最终架构：DeepTutor 负责课程结构管理与前端交互，Jellyfin 负责视频存储与播放引擎，"
       "二者通过后端 API 打通。原 DeepTutor 的 AI 功能全部保留，课程点播是叠加在其上的新能力。")

doc.add_page_break()

# ═══════════════════════ 系统架构 ═══════════════════════
h1("系统架构图")
para("整体架构如下图所示：前端（Next.js）通过 REST API 访问 DeepTutor 后端（FastAPI），"
     "后端负责课程管理、认证与视频流代理；Jellyfin 作为媒体引擎负责视频入库与转码；"
     "数据存储在 SQLite（课程/视频/进度元数据）与磁盘（视频文件）两层。", indent=True)

# 精确版架构图（正式）
try:
    pic_path = r"D:\Users\zzh\Desktop\DeepTutor项目\架构图_精确版.png"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(pic_path, width=Cm(15.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run("图 1：DeepTutor 课程点播系统架构图（精确版）")
    set_cn(r, "宋体", 9, False, RGBColor(0x64, 0x74, 0x8B))
except Exception as e:
    para(f"（架构图插入失败: {e}）", indent=False)

doc.add_page_break()

# ═══════════════════════ 二、已完成工作 ═══════════════════════
h1("二、已完成工作")

h2("2.1 阶段一：技术验证（P0）")
para("在写任何业务代码之前，我先做了完整的 P0 验证，确认技术路线可行：")
bullet("验证了 Jellyfin API 的连通性与库结构识别能力；")
bullet("验证了课程视频按 SxxExx 剧集命名能被 Jellyfin 正确识别为「一集」；")
bullet("调整了媒体库路径配置，避免各库之间互相抢占路径；")
bullet("发现并规避了一个关键陷阱：Jellyfin 的删除接口会级联删除磁盘上的媒体文件，"
       "后续所有删除操作都做了防护。")

h2("2.2 阶段二：后端能力（P1+P2）")
para("搭建了完整的课程媒体后端：")
bullet("Jellyfin 客户端封装（入库、扫描、播放信息、流代理）；")
bullet("课程/视频的 SQLite 存储层（课程、视频、进度三张表）；")
bullet("10 个 REST 端点：课程增删查、视频上传/删除/播放、进度保存/读取；")
bullet("上传管线：文件落盘 → 触发 Jellyfin 扫描 → 轮询确认索引成功，全自动完成。")

h2("2.3 阶段三：播放能力（P3）")
bullet("实现了流代理：前端 <video> 标签通过我们的后端代理 Jellyfin 视频流，"
       "Jellyfin 的 API Key 永远不会暴露给浏览器；")
bullet("视频播放支持拖拽进度（Range 请求透传）；")
bullet("新增了自动转码能力：非 H.264 格式（如 MKV/HEVC）会自动转码成浏览器可播的 H.264/AAC/MP4，"
       "H.264 原生格式则直接传输，零 CPU 开销。")

h2("2.4 阶段四：用户与权限（P4）")
bullet("完整的注册/登录体系（JWT Cookie，24 小时有效）；")
bullet("角色隔离：管理员可管理课程和用户，普通学生只能观看；")
bullet("第一个注册用户自动成为管理员，之后由管理员在后台创建学生账号。")

h2("2.5 阶段五：部署与生产化（P5）")
bullet("编写了 docker-compose.media.yml，一条命令拉起 DeepTutor + Jellyfin 全家桶；")
bullet("编写了完整的部署文档（架构图、部署步骤、HTTPS 配置、备份方案、故障排查表）；")
bullet("完成了局域网部署：前后端都绑定 0.0.0.0，其他设备可通过局域网 IP 直接访问。")

doc.add_page_break()

# ═══════════════════════ 三、界面与体验优化 ═══════════════════════
h1("三、界面与体验优化")
para("在核心功能全部跑通之后，我以产品视角对课程模块做了三轮体验优化，"
     "每轮都由 AI 产品评审（pi）独立评估、指出问题、我再修正，确保不是自我感觉良好。")

h2("3.1 第一轮（P0）：播放器体验")
make_table(
    ["改进项", "说明"],
    [
        ["键盘快捷键", "空格暂停、J/L 快退快进 10 秒、方向键微调、M 静音、F 全屏、逗号句号调速"],
        ["播放进度记忆", "播放位置自动保存，刷新/换设备后从上次位置继续（双端存储）"],
        ["上/下一集 + 自动续播", "播放器下方按钮切换，一集看完自动播下一集"],
        ["倍速控制", "1x / 1.25x / 1.5x / 2x 一键切换（学生学习刚需）"],
        ["上传进度条 + 文件校验", "大文件上传实时百分比，限制格式（MP4/WebM/MKV 等）与大小（2GB）"],
        ["删除二次确认", "删除课程/视频前弹窗确认，防误删"],
    ],
    widths=[4.0, 10.5],
)

h2("3.2 第二轮（P1）：学习进度反馈")
make_table(
    ["改进项", "说明"],
    [
        ["已观看标记", "播完 80% 的视频在课时列表打绿色对勾"],
        ["课程总体进度", "列表卡片显示「已完成 X/Y 课时」+ 绿色进度条"],
        ["继续上次学习", "进入课程列表页顶部显示最近观看的课程与时间点，一键续学"],
        ["播放错误重试", "视频加载失败显示明确错误 + 重试按钮"],
        ["加载骨架屏", "课程列表加载时显示卡片占位，避免内容跳动"],
        ["移动端抽屉导航", "手机上侧边栏改为汉堡菜单 + 滑出抽屉，不再挤压内容"],
    ],
    widths=[4.0, 10.5],
)

h2("3.3 第三轮（P2）：跨设备与国际化")
make_table(
    ["改进项", "说明"],
    [
        ["服务端进度持久化", "播放进度同步到服务器（按用户隔离），换设备不丢进度——发布阻塞项"],
        ["离线进度回补", "断网期间本地记的进度，联网后自动补传服务器"],
        ["课程搜索", "课程列表页按标题/简介实时过滤"],
        ["语言统一", "登录/注册等残留英文全部改为中文"],
        ["深色模式", "跟随系统偏好，侧边栏一键切换"],
    ],
    widths=[4.0, 10.5],
)

doc.add_page_break()

# ═══════════════════════ 四、质量保障 ═══════════════════════
h1("四、质量保障")
bullet("新增回归测试（tests/test_course_progress.py）：覆盖进度保存/读取、用户隔离、覆盖更新、"
       "课程聚合、外键级联删除，5 个用例全部通过；")
bullet("SQLite 开启外键约束，删除课程/视频时进度数据自动级联清理，不留孤儿数据；")
bullet("每轮改动均通过 AI 评审 + 全页面 HTTP 200 冒烟验证；")
bullet("所有代码已提交并推送 GitHub 远程仓库（36 个提交），版本可回退。")

# ═══════════════════════ 五、当前状态 ═══════════════════════
h1("五、当前状态")
make_table(
    ["服务", "地址", "说明"],
    [
        ["前端", "http://localhost:3000（LAN: http://100.80.148.8:3000）", "Windows 原生进程，绑定 0.0.0.0"],
        ["后端 API", "http://localhost:8001", "FastAPI + uvicorn，绑定 0.0.0.0"],
        ["Jellyfin", "localhost:8096", "Docker 容器，仅后端访问（不直接暴露）"],
        ["测试账号", "admin / admin12345（管理员）；student1 / student123（学生）", "可登录体验"],
    ],
    widths=[3.0, 6.5, 5.0],
)
para("目前已上传「Python 入门」课程（4 个课时）作为演示数据，完整走通了 上传 → 索引 → 播放 全流程，"
     "包括转码路径（HEVC 测试视频已实测可播放）。")

# ═══════════════════════ 六、下一步计划 ═══════════════════════
h1("六、下一步计划")
make_table(
    ["事项", "说明", "优先级"],
    [
        ["小范围内测", "邀请 10-50 名学生试用，收集真实反馈", "立即"],
        ["笔记/书签功能", "视频时间戳标注，学生边看边记（pi 评审建议）", "内测后"],
        ["章节目录/时间戳跳转", "长视频章节导航", "内测后"],
        ["错误监控与埋点", "Sentry 上报 + 观看行为统计（完课率、流失点）", "内测后"],
        ["HTTPS 与公网部署", "Caddy/Nginx 反代，支持校外访问（文档已备）", "按需"],
    ],
    widths=[4.5, 8.0, 2.0],
)

# ═══════════════════════ 七、风险与说明 ═══════════════════════
h1("七、需要说明的事项")
bullet("技术路线与原计划有一处调整：原计划融合 seerr，经验证改为融合 Jellyfin。"
       "原因如上文所述（seerr 依赖 TMDB 影视元数据，无法支持自制课程）。"
       "如果领导坚持要保留 seerr，后续可以加一层适配，但建议以 Jellyfin 为主；")
bullet("当前转码为纯 CPU 转码（服务器无 GPU），多路并发转码时会有性能瓶颈，"
       "内测阶段影响可控，正式上线前可评估加显卡或改用 H.264 源文件；")
bullet("局域网 IP（100.80.148.8）会随网络环境变化，跨网络使用需配置固定 IP 或域名。")

doc.save(r"D:\Users\zzh\Desktop\DeepTutor项目\DeepTutor课程点播系统-工作总结-v2.docx")
print("文档已生成")
